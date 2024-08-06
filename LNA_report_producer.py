# -*- encoding: utf-8 -*-
"""
This file produces reports of the LNAs performances of Band 2 ALMA-ECOGAL experiment.

@author: Francesco Andreetto

Created on: May 23rd 2024, European Southern Observatory (ESO), Garching Bei München (Germany)
Last editing: August 5th 2024, Observatory of Astrophysics and Space Science (INAF-OAS) - Bologna (Italy)
"""

# Modules & Libraries

import argparse
import logging
from docx import Document
from docx.shared import Pt, RGBColor

import f_Report as fr
from pathlib import Path
from rich.logging import RichHandler

# Use the module logging to produce nice messages on the shell
logging.basicConfig(level="INFO", format='%(message)s',
                    datefmt="[%X]", handlers=[RichHandler()])


def main():
    # ==================================================================================================================
    # Reports Requirements

    logging.info('\nLoading dir and templates information...')

    # Command Line used to start the pipeline
    # command_line = " ".join(sys.argv)

    # Create the argument parser by instantiating ArgumentParser
    # Note: the description is optional. It will appear if the help (-h) is required from the command line
    parser = argparse.ArgumentParser(prog='PROGRAM', description="ECOGAL-Band 2 Script to parse "
                                                                 "the Gain and Noise T conditions of the LNAs.")

    ####################################################################################################################
    # Flags of the CLI
    ####################################################################################################################
    # path_file
    parser.add_argument("--path", action='store', type=str,
                        help='- Location of the txt files used to create plots (database)')
    # report_dir
    parser.add_argument("--report_dir", action='store', type=str, default="/REPORTS",
                        help='- Directory where to save all the reports')

    # font_name
    parser.add_argument("--font_name", action='store', type=str, default="Times New Roman",
                        help='- Font of the text in the document')

    # Call .parse_args() on the parser to get the Namespace object that contains all the user’s arguments.
    args = parser.parse_args()
    logging.info(args)

    report_dir = "REPORTS"
    # Check if the dir exists. If not, it will be created.
    Path(report_dir).mkdir(parents=True, exist_ok=True)
    # ==================================================================================================================

    # ==================================================================================================================
    # REPORT Production
    # ==================================================================================================================

    logging.info(f"\nI am putting the report into: {report_dir}.\n\n")

    # Operations
    # ==================================================================================================================

    # Document name
    report_file_name = "LNA_Report.docx"
    # Document final location
    output_path = f"{report_dir}/{report_file_name}"

    # Duplicate the template
    fr.duplicate_file(input_file=f"{report_dir}/Band 2 Cryogenic LNAs Delivery Report Production Batch 1_template.docx",
                      output_file=output_path)
    logging.info(f"Template duplicated and renamed successfully: {output_path}.\n")

    # Define the document to use
    doc = Document(output_path)

    # Define document classification
    # classification = "ESO Internal Use [Confidential for Non-ESO Staff]"
    # core_properties = doc.core_properties
    # core_properties.subject = classification

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # PAG 1 - GIVEN BY THE TEMPLATE ON TOP
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================

    logging.info("Template header completed\n"
                 "==================================================================\n")

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # PAG 2
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================

    # ------------------------------------------------------------------------------------------------------------------
    # AUTHOR TABLE
    # ------------------------------------------------------------------------------------------------------------------
    # Load information from the file
    authors_dict = fr.load_dict_from_datafile(file_path="Tables/authors.txt")
    # Fill the Author Table
    fr.fill_table(doc=doc, data=authors_dict, table_index=2)

    # ------------------------------------------------------------------------------------------------------------------
    # CHANGES TABLE
    # ------------------------------------------------------------------------------------------------------------------
    # Load information from the file
    changes_dict = fr.load_dict_from_datafile(file_path="Tables/change_record.txt")
    # Fill the Changes Table
    fr.fill_table(doc=doc, data=changes_dict, table_index=3)
    # ------------------------------------------------------------------------------------------------------------------

    # ==================================================================================================================
    logging.info("Authors & Changes Table filled in the template.\nStarting sections.\n"
                 "==================================================================\n")

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # SECTION 1 - Description
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================

    # Add Section 1 title
    fr.write_nice_heading(doc=doc, text="Description", level=1,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=18, font_style="bold")
    # ------------------------------------------------------------------------------------------------------------------
    # Add Subsection 1.1 title
    fr.write_nice_heading(doc=doc, text="Purpose", level=2,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=16, font_style="normal")
    # Read Subsection 1.1
    text = fr.read_text_from_file("Sections/pag4_sec1_1.txt")
    # Write Subsection 1.1
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # ------------------------------------------------------------------------------------------------------------------
    # Add Subsection 1.2 title
    fr.write_nice_heading(doc=doc, text="Scope", level=2,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=16, font_style="normal")
    # Read Subsection 1.2
    text = fr.read_text_from_file("Sections/pag4_sec1_2.txt")
    # Write Subsection 1.2
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # ------------------------------------------------------------------------------------------------------------------
    # Add Subsection 1.3 title
    fr.write_nice_heading(doc=doc, text="Reference Documents List", level=2,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=16, font_style="normal")
    # Reference Documents List
    ref_dict = fr.load_dict_from_datafile("Tables/reference_document_list.txt")
    # Make the Reference Document List Table
    fr.make_table(doc=doc, data=ref_dict, title="\n")

    logging.info("Section 1 completed\n"
                 "==================================================================\n")
    # Go to a new page
    doc.add_page_break()

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # SECTION 2 - Acceptance and cascading CLNAs
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================

    # Add Section 2 title
    fr.write_nice_heading(doc=doc, text="Acceptance and cascading CLNAs", level=1,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=18, font_style="bold")
    # Read Section 2
    text = fr.read_text_from_file("Sections/pag5_sec2.txt")
    # Write Section 2
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # Add Subsection 2.1 title
    fr.write_nice_heading(doc=doc, text="Individual CLNAs waivers", level=2,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=16, font_style="normal")

    # Write the CLNA_waivers.txt file
    fr.create_waivers_file(input_file_path="LNA_couples/LNA_couples.txt",
                           directory_path_MPI="LNA_specifications/MPI_specifications/Waivers",
                           directory_path_LNF="LNA_specifications/LNF_specifications/Waivers")

    # Individual CLNAs Documents TABLE
    document_dict = fr.load_dict_from_datafile("Tables/CLNA_waivers.txt")
    # Make the Reference Document List Table
    fr.make_table(doc=doc, data=document_dict, title=" ")

    logging.info("Section 2 completed\n"
                 "==================================================================\n")
    # Go to a new page
    doc.add_page_break()

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # SECTION 3 - Test system at ESO
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================

    # Add Section 3 title
    fr.write_nice_heading(doc=doc, text="Test system at ESO", level=1,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=18, font_style="bold")
    # Read Section 3
    text = fr.read_text_from_file("Sections/pag7_sec3.txt")
    # Write Section 3
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # Figure 3.1
    img_path = "Figures/Figure_3_1.png"
    caption_file = "Figures/Caption_figure_3_1.txt"
    fr.add_figure(doc=doc, img_path=img_path, caption_file=caption_file, img_width=300, img_height=200)

    # Read Section 3 pt 2
    text = fr.read_text_from_file("Sections/pag7_sec3_pt2.txt")
    # Write Section 3 pt2
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # Figure 3.2
    img_path = "Figures/Figure_3_2.png"
    caption_file = "Figures/Caption_figure_3_2.txt"
    fr.add_figure(doc=doc, img_path=img_path, caption_file=caption_file, img_width=250)

    # Read Section 3 pt 3
    text = fr.read_text_from_file("Sections/pag7_sec3_pt3.txt")
    # Write Section 3 pt3
    fr.write_nice_text(doc=doc, text=text, output_path=output_path)

    # Figure 3.3
    img_path = "Figures/Figure_3_3.png"
    caption_file = "Figures/Caption_figure_3_3.txt"
    fr.add_figure(doc=doc, img_path=img_path, caption_file=caption_file, img_width=300, img_height=200)

    logging.info("Section 3 completed\n"
                 "==================================================================\n")
    # Go to a new page
    doc.add_page_break()

    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # SECTION 4 and following
    # ==================================================================================================================
    # ==================================================================================================================
    # ==================================================================================================================
    # Section Counter
    sect_counter = 0

    plot_path = "LNA_Plots/"
    # Get plot names from the directory
    plot_names = fr.get_plot(plot_path)

    for idx, item in enumerate(plot_names):

        # Set margins of the section
        fr.set_margins(doc, top=1, bottom=1, left=0.7, right=0.7)

        # fr.write_chain_section(item)
        # Get the name of the first LNA
        name_LNA_1 = item[-26:-17]

        # Get the name of the second LNA
        name_LNA_2 = item[-16:-8]

        # Write the heading of the section
        text = f"CLNA chain: {name_LNA_1} and {name_LNA_2}"
        sect_counter = idx + 4
        fr.write_nice_heading(doc=doc, text=text, level=1,
                              font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                              font_size=18, font_style="italic")
        # Write some text in the section
        text = f"This CLNA chain is configured as {name_LNA_1} (1st stage), and {name_LNA_2} (2nd stage)."
        fr.write_nice_text(doc=doc, text=text, output_path=output_path,
                           font_name='Times New Roman', font_size=10, font_color=(0, 0, 0), font_style=None)

        # --------------------------------------------------------------------------------------------------------------
        # MAKE BIAS TABLES

        # Add some white space before the table
        doc.add_paragraph(' ')

        # BIAS TABLE 1 - Set Values
        #  =============================================================================================================
        logging.info("Creating table: recommended bias (SET VALUES).\n")
        table_title = (f"Table {idx + 4}.1 - Bias conditions - Set Values.\n "
                       f"Best values recommended by ESO, for 15K ambient operation.")

        # Directory of the bias txt files
        dir_path = "Bias_Txt_Files"
        # Get the filename needed
        filename = fr.get_file_from_string(directory_path=dir_path, search_string=name_LNA_1, n_char=-1)
        # Load the dictionary
        data = fr.load_dict_from_filename(filename_path=f"{dir_path}/{filename}")
        # Write the set-values on the table in the doc
        fr.make_table(doc=doc, data=data, title=table_title)
        #  =============================================================================================================

        # Add some white space between the two tables
        doc.add_paragraph(' ')

        # BIAS TABLE 2 - Read out Values
        #  =============================================================================================================
        # Extract bias read out values from the datafile
        logging.info(f"Extracting bias read out values for {name_LNA_1} and {name_LNA_2}\n")
        fr.extract_lna_bias_values(directory_path="Bias_Config_Files", lna1_name=name_LNA_1, lna2_name=name_LNA_2)

        logging.debug(f"1): {name_LNA_1} \n2): {name_LNA_2}")

        logging.info("Creating table: first measure bias (READ OUT VALUES).\n")
        table_title = (f"Table {idx + 4}.2 - Bias conditions - Read Out Values.\n"
                       f"Bias values used to start the fine-tuning measurements at cryogenic conditions.")
        # Load the dictionary reading the configuration files
        data = fr.load_dict_from_datafile(f"Tables/bias_table_{name_LNA_1}_{name_LNA_2}.txt")
        # Write the read-values on the table in the doc
        fr.make_table(doc=doc, data=data, title=table_title)

        # Add some white space before the figure
        doc.add_paragraph(' ')

        # Add Plot of the LNA
        fr.add_figure(doc=doc, img_path=f"LNA_Plots/{item}", caption_file="none", img_width=350, img_height=250)

        # Add Individual CLNA info
        text = "Individual CLNAs, test reports:"
        fr.write_nice_heading(doc=doc, text=text, level=0,
                              font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                              font_size=11, font_style="bold")

        # ==============================================================================================================
        # Test Reports
        # ==============================================================================================================
        # 1. Compliance Info
        # --------------------------------------------------------------------------------------------------------------
        lna_1_info = fr.get_file_from_string(directory_path="LNA_specifications/MPI_specifications",
                                             search_string=f"{name_LNA_1.replace('.', ' ')}",
                                             n_char=19)
        text = f"Compliance verification for amplifier: {name_LNA_1}\t\t{lna_1_info}"
        fr.write_nice_text(doc=doc, text=text, output_path=output_path, font_size=8)

        # 2. Test data info
        # --------------------------------------------------------------------------------------------------------------
        modified_name_LNA_2 = name_LNA_2[-4:]
        logging.debug(modified_name_LNA_2)

        # Get pdf file name
        pdf_test_file_name = fr.get_file_from_string(directory_path="LNA_specifications/LNF_specifications",
                                                     search_string=modified_name_LNA_2,
                                                     n_char="all")
        # Read the first line of the Test data file
        logging.debug(pdf_test_file_name)
        lines = fr.get_lines_from_pdf(pdf_file_path=f"LNA_specifications/LNF_specifications/{pdf_test_file_name}")
        # logging.debug(lna_2_info)
        if not lines:
            logging.debug("No PDF file found, hence no reference file in the Report.")
            text = " "
        else:
            logging.debug(f"{lines[0]}\t{lines[1]}")
            text = f"{lines[0]}\t\t{lines[1]}"

        fr.write_nice_text(doc=doc, text=text, output_path=output_path, font_size=8)

        # --------------------------------------------------------------------------------------------------------------
        # Applicable waiver
        # --------------------------------------------------------------------------------------------------------------

        text = "Individual CLNAs, applicable waivers:"
        fr.write_nice_heading(doc=doc, text=text, level=0,
                              font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                              font_size=11, font_style="bold")

        logging.debug(name_LNA_1[-3:])
        logging.debug(name_LNA_2[-4:])

        # line_1 = fr.find_keyword_in_file(file_path="Tables/individual_CLNA_waivers.txt", keyword=name_LNA_1)
        line_1 = fr.find_keyword_in_file(file_path="Tables/CLNA_waivers.txt", keyword=name_LNA_1[-3:])
        logging.debug(line_1)

        # line_2 = fr.find_keyword_in_file(file_path="Tables/individual_CLNA_waivers.txt", keyword=name_LNA_2)
        line_2 = fr.find_keyword_in_file(file_path="Tables/CLNA_waivers.txt", keyword=name_LNA_2[-4:])
        logging.debug(line_2)

        # Collect & write the text in the report
        text = f"{line_1}\n{line_2}"
        fr.write_nice_text(doc=doc, text=text, output_path=output_path, font_size=8)

        # Go to a new page
        doc.add_page_break()
        logging.info(f"Section {4 + idx} completed!\n"
                     f"==================================================================\n")

    # Write the heading of the section
    sect_counter += 1
    text = f"Summary of all CLNAs performance"
    fr.write_nice_heading(doc=doc, text=text, level=1,
                          font_color=RGBColor(0, 0, 0), font_name="Times New Roman",
                          font_size=18, font_style="italic")
    # Add Plot of the LNA
    fr.add_figure(doc=doc, img_path=f"LNA_Plots/ALL/Plot_All.png",
                  caption_file="none", img_width=450, img_height=450)
    # Move the "End of document" text at the very end of the document
    fr.move_text_to_end(doc=doc, search_text="--- End of document ---", output_path=f"{output_path}")

    ####################################################################################################################
    # Save the document
    doc.save(f'{output_path}')
    logging.info(f"File Saved: {output_path}")
    ####################################################################################################################


if __name__ == '__main__':
    main()
