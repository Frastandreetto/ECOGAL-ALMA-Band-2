# -*- encoding: utf-8 -*-
"""
This file contains the main functions used in ALMA-ECOGAL project producing reports for tests of LNAs @ ESO.

@author: Francesco Andreetto

Created on: May 29th 2024, European Southern Observatory (ESO) - Garching Bei München (Germany)
Last editing: August 5th 2024, Observatory of Astrophysics and Space Science (INAF-OAS) - Bologna (Italy)

"""

# Import Modules
import PyPDF2
import logging
import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_figure(doc: Document, img_path: str, caption_file: str, img_width=400, img_height=400,
               font_name='Times New Roman', font_size=10, font_color=(0, 0, 0)):
    """
    Add a centered figure with a caption to the docx file.

    Parameters:\n
    - **doc** (``Document``): docx in which put the image;
    - **img_path** (``str``): path of the image;
    - **caption_file** (``str``): path of the file txt with the caption of the image;
    - **img_width** (``int``): width of the image;
    - **img_height** (``int``): height of the image;
    - **font_name** (``str``): name of the font of the paragraph;
    - **font_size** (``int``): size of the title (Pt);
    - **font_color** (``RGBColor``): color of the font.
    """
    # Creating a paragraph
    paragraph = doc.add_paragraph()
    # Center the paragraph with the image
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    # Add the picture
    run.add_picture(img_path, width=Pt(img_width), height=Pt(img_height))

    if caption_file != "none":
        # Read the caption from the caption file
        caption = read_text_from_file(caption_file)

        # Add the caption to the doc
        run = paragraph.add_run(caption)
        # Center the caption
        run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Font of the text
        run.font.name = font_name
        run.font.size = Pt(font_size)

        # Color of the font
        run.font.color.rgb = RGBColor(*font_color)


def create_waivers_file(input_file_path: str, directory_path_MPI: str, directory_path_LNF: str,
                        output_file_name='Tables/CLNA_waivers.txt'):
    """
    Read from an input file the names of the LNAs used in the analysis, then look for waiver files into two directories,
    one per LNA-type of the chain, and create an output files in which writes the names of these files, if found.

    Parameters:\n
    - **input_file_path** (``str``): path of the input file in which are written the names of the LNAs analyzed;
    - **directory_path_MPI** (``str``): path of the directory containing the MPI waiver files;
    - **directory_path_LNF** (``str``): path of the directory containing the LNF waiver files;
    - **output_file_name** (``str``): name of the output_file containing all the names of the waiver files.
    """

    # Initialize the counter for the reference documents
    ref_counter = 0

    try:
        # Read the input file
        with open(input_file_path, 'r') as input_file:
            lna_couples = input_file.readlines()

        # Open the output file for writing
        with open(output_file_name, 'w') as output_file:
            # Write the header
            output_file.write("'Reference' 'Document title and number' 'Status'\n")

            # Process each LNA couple and write the corresponding output
            for i, lna in enumerate(lna_couples):
                lna = lna.strip()
                if not lna:
                    logging.warning(f"No LNA name found: {lna}")
                    continue

                # Odd lines - MPI LNAs
                if i % 2 == 0:
                    # Get the file name from the serial number Wx.XX-XXX
                    filename = get_file_from_string(directory_path_MPI, lna[-6:], n_char="all")
                    if filename == "":
                        pass
                    else:
                        logging.warning(f"File name: {filename}.\n")
                        pdf_path = os.path.join(directory_path_MPI, filename)
                        logging.debug(f"{pdf_path}\n\n")

                        # Get the first two line of the file
                        lines = get_lines_from_pdf(pdf_file_path=pdf_path)
                        text = f"{lines[6]} {lines[7]} {lines[8]}"

                # Even lines - LNF LNAs
                else:
                    # Get the file name from the serial number LNF-XXXX
                    logging.debug(lna[-4:])
                    filename = get_file_from_string(directory_path_LNF, f"waiver_{lna[-4:]}", n_char="all")
                    if filename == "":
                        pass
                    else:
                        logging.warning(f"File name: {filename}.\n")
                        pdf_path = os.path.join(directory_path_LNF, filename)

                        # Get the first two line of the file
                        lines = get_lines_from_pdf(pdf_file_path=pdf_path)
                        text = f"{lines[8]} {lines[9]}"

                if filename == "":
                    pass
                else:
                    # Increasing the ref counter
                    ref_counter += 1
                    # Fill the line with the ref info
                    reference = f"[RW{ref_counter}]"
                    status = "approved"
                    output_file.write(f"'{reference}' '{text}' '{status}'\n")

        logging.info(f"Output file '{output_file_name}' created successfully.\n")

    except Exception as e:
        logging.error(f"Waiver Error: {e}")


def duplicate_file(input_file: str, output_file: str):
    """
    Duplicates a file and renames it.

    Parameters:
    - **input_file** (``str``): path to the source file;
    - **output_file** (``str``): path to the destination file.
    """
    try:
        # Open the input file in binary read mode
        with open(input_file, 'rb') as src_file:
            # Read the content of the input file
            content = src_file.read()

        # Open the output file in binary write mode
        with open(output_file, 'wb') as dest_file:
            # Write the content to the output file
            dest_file.write(content)

        logging.info(f"File duplicated and renamed from '{input_file}' to '{output_file}' successfully.\n")
    except Exception as e:
        print(f"Duplication Error: {e}")


def extract_lna_bias_values(directory_path: str, lna1_name: str, lna2_name: str):
    """
    Extracts bias values for two LNAs from a txt file in the specified directory.
    Writes the extracted values into an output file in the 'Tables' directory (one level above the ``directory_path``)
    The output file is named in the format 'bias_table_'``lna1_name``'_'``lna2_name``'.txt', and contains a header line
    followed by the extracted values for both LNAs in a structured format.

    Parameters:\n
    - **directory_path** (``str``): path to the directory where the input file is located;
    - **lna1_name** (``str``): name of the first LNA to be used in the file search and value extraction;
    - **lna2_name** (``str``): name of the second LNA to be used in the file search and value extraction.
    """

    # Find the input file in the directory
    input_file = None
    for filename in os.listdir(directory_path):
        if lna1_name in filename and lna2_name in filename:
            input_file = os.path.join(directory_path, filename)
            break

    if input_file is None:
        logging.warning(f"No file found containing {lna1_name} and {lna2_name} in the name.\n")
        return

    # Read the input file
    with open(input_file, 'r') as file:
        lines = file.readlines()

    # Extracting the values for each LNA
    lna1_values = []
    lna2_values = []
    lna1_prefix = "LNA_1"
    lna2_prefix = "LNA_2"

    for line in lines:
        if lna1_prefix in line:
            value = line.split(':')[1].strip().replace(' V', '').replace(' mA', '')
            lna1_values.append(value)
        elif lna2_prefix in line:
            value = line.split(':')[1].strip().replace(' V', '').replace(' mA', '')
            lna2_values.append(value)

    # Create the output content
    header = "'LNA name' 'Vd1 [V]' 'Id1 [mA]' 'Vg1 [V]' 'Vd2 [V]' 'Id2 [mA]' 'Vg2 [V]' 'Vd3 [V]' 'Id3 [mA]' 'Vg3 [V]'"
    lna1_line = (f"'{lna1_name}' "
                 f" '{lna1_values[0]}' '{lna1_values[1]}' '{lna1_values[2]}'"
                 f" '{lna1_values[3]}' '{lna1_values[4]}' '{lna1_values[5]}'"
                 f" '{lna1_values[6]}' '{lna1_values[7]}' '{lna1_values[8]}'")
    lna2_line = (f"'{lna2_name}' "
                 f"'{lna2_values[0]}' '{lna2_values[1]}' '{lna2_values[2]}'"
                 f" '{lna2_values[3]}' '{lna2_values[4]}' '{lna2_values[5]}'"
                 f" '{lna2_values[6]}' '{lna2_values[7]}' '{lna2_values[8]}'")

    output_content = f"{header}\n{lna1_line}\n{lna2_line}\n"

    # Write the output content to a new file
    output_file = os.path.join(f"{directory_path}/../Tables/", f"bias_table_{lna1_name}_{lna2_name}.txt")
    with open(output_file, 'w') as file:
        file.write(output_content)


def fill_table(doc: Document, data: dict, table_index: int):
    """
    Fills a specified table in the provided Document object with the provided data.
    Headers in the table are set in bold.

    Parameters:\n
    - **doc** (``Document``): Document object representing the Word document;
    - **data** (``dict``): dict where keys are column headers and values are lists of data to fill the columns;
    - **table_index** (``int``): index of the table to be filled in the doc.
    """
    # Get the specified table in the document
    tables = doc.tables
    if not tables:
        raise ValueError("No tables found in the document.")

    if table_index >= len(tables):
        raise IndexError("Table index out of range.")

    table = tables[table_index]

    # Ensure the table has the correct number of columns
    num_cols = len(data)
    if len(table.columns) < num_cols:
        raise ValueError("Table does not have enough columns for the data provided.")

    # Fill in headers with bold formatting
    headers = list(data.keys())
    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        cell.text = header
        # Set the header to bold
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Fill in data rows
    for col_index, (header, values) in enumerate(data.items()):
        for row_index, value in enumerate(values):
            # Check if the row exists
            if row_index + 1 < len(table.rows):
                table.cell(row_index + 1, col_index).text = value
            else:
                # Add new row if it doesn't exist
                row = table.add_row()
                row.cells[col_index].text = value


def find_chain(file_path: str) -> (str, str):
    """
    Find the names of the LNAs in a chain from the datafile.
    Parameters:\n
    - **file_path** (``str``): path of the file to read.
    """
    for filename in os.listdir(file_path):
        if filename.startswith("20") and filename.endswith(".txt"):
            # Return the names of the two LNAs
            return filename[11:19], filename[21:-4]


def find_keyword_in_file(file_path: str, keyword: str) -> str:
    """
    Search for a keyword in a given file and returns the line containing the keyword + the reference [RW]

    Parameters:\n
    - **file_path** (``str``): path to the text file;
    - **keyword** (``str``): keyword to search for in the file.

    Returns:\n
    - ``str``: the cleaned line containing the keyword and the corresponding reference [RW];
    - ``---``: if the keyword is not found in the file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                if keyword in line:
                    # Remove single quotes and the word 'approved'
                    cleaned_line = line.replace("'", "").replace("approved", "").strip()
                    return cleaned_line
        return "---"
    except Exception as e:
        print(f"Keyword Error: {e}")
        return "---"


def get_file_from_string(directory_path, search_string, n_char=19) -> str:
    """
    Searches for a file within the specified directory and return its name if it contains the given search string.

    Parameters:\n
    - **directory_path** (``str``): path to the directory where the files are stored;
    - **search_string** (``str``): string to search for within filenames in the directory;
    - **n_char** (``int`` or ``str``): The number of characters to return from the start of the filename.
      If set to "all", the entire filename is returned. Defaults: 19.

    Returns:\n
    - **str**: The first `n_char` characters of the filename containing the `search_string`.
      If `n_char` is "all", the entire filename is returned. If no file is found, an empty string is returned.
    """
    # Initialize a str to store results
    results = ""

    # Iterate through all files in the directory
    for filename in os.listdir(directory_path):
        # Check if the search string is in the filename
        if search_string in filename:
            # Append the first n_char characters of the filename to the results list
            if n_char == "all":
                results = filename
            else:
                results = filename[:n_char]

    return results


def get_lines_from_pdf(pdf_file_path: str) -> list:
    """
     Extracts and returns all the lines of the PDF file specified by ``pdf_file_path`` as a list.

    Parameters:\n
    - **pdf_file_path** (``str``): The path to the PDF file from which lines are to be extracted.

    Returns:\n
    - **list**: A list of strings, each representing a line of text extracted from the PDF.
    If there are no lines in the PDF, an error occurs and the function returns a list like this: ["---", "---"].
    """
    try:
        # Open the PDF file
        with open(pdf_file_path, 'rb') as file:
            reader = PyPDF2.PdfFileReader(file)
            all_lines = []

            # Iterate through all the pages in the PDF
            for page_num in range(reader.numPages):
                # Extract text from each page
                page_text = reader.getPage(page_num).extract_text()
                if page_text:
                    # Split text into lines and add them to the list
                    lines = page_text.split('\n')
                    all_lines.extend([line.strip() for line in lines if line.strip()])

            return all_lines

    except Exception as e:
        logging.error(f"PDF lines Error: {e}")
        return ["---", "---"]


def get_plot(dir_path: str) -> list:
    """
    Get the plots from a given directory and store the names into a list.
    Parameters:
    - **dir_path** (``str``): path of the directory that contains the plots.
    """
    # List to store the names of image files
    image_names = []

    # Iterate over all files in the specified directory
    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)

        # Check if it is a file and ends with .png
        if os.path.isfile(file_path) and filename.lower().endswith('.png'):
            image_names.append(filename)

    return image_names


def load_dict_from_datafile(file_path: str) -> dict:
    """
    Load a dictionary from a file of data.
    Parameters:\n
    - **file_path** (``str``): path of the file to read.
    Returns:\n
    - **data_dict** (``dict``): dict filled with the data of the file.
    """
    data_dict = {}
    keys = None

    with open(file_path, 'r') as file:
        for idx, line in enumerate(file):
            line = re.findall(r"'(.*?)'", line)

            if idx == 0:
                # Save the headers in the dict
                keys = line
                for key in keys:
                    # Initialize an empty list linked to the key
                    data_dict[key] = []
            else:
                # Save the values in the dict
                for i, value in enumerate(line):
                    key = keys[i]
                    data_dict[key].append(value)

    return data_dict


def estimate_text_width(text, font_size):
    """
    Estimate the width of text in inches based on character count and font size.
    """
    avg_char_width = 0.5  # Average character width in inches (approximation)
    text_length = len(text)
    return avg_char_width * text_length * (font_size / 8)


def move_text_to_end(doc: Document, search_text: str, output_path: str):
    """
    Move the paragraph containing the ``search_text`` to the end of the document and save to the specified path.

    Parameters:
    - **doc** (``Document``): The Document object;
    - **search_text** (``str``): The text to search for and move;
    - **output_path** (``str``): Path to save the modified document.
    """

    # Find the paragraph that contains the search text
    para_to_move = None
    for para in doc.paragraphs:
        if search_text in para.text:
            para_to_move = para
            break

    if para_to_move is None:
        raise ValueError(f"Paragraph containing '{search_text}' not found.")

    # Access the underlying XML of the document
    doc_xml = doc._element
    body = doc_xml.find('.//w:body', namespaces=doc_xml.nsmap)

    # Remove the paragraph from its current position
    para_element = para_to_move._element
    body.remove(para_element)

    # Append the paragraph to the end of the document
    body.append(para_element)

    # Save the modified document
    doc.save(output_path)


def make_table(doc: Document, data: dict, title: str, header_font_size=10, elements_font_size=7):
    """
    Create a table on a docx from a dictionary.
    The keys of the dictionary will be the headers of the table, written in bold.
    """
    # Add a title for the table
    write_nice_heading(doc=doc, text=f"{title}", level=0,
                       font_color=RGBColor(0, 0, 0), font_name='Times New Roman',
                       font_size=12, font_style="italic")

    # Create a table with number of rows and columns (+1 for the heading)
    num_rows = len(next(iter(data.values()))) + 1
    num_col = len(data)
    new_table = doc.add_table(rows=num_rows, cols=num_col)

    # Add the heading to the table
    head = new_table.rows[0].cells
    max_widths = []

    for i, col in enumerate(data.keys()):
        # Set the headers to the columns
        head[i].text = col
        max_width = estimate_text_width(col, header_font_size)
        max_widths.append(max_width)
        for run in head[i].paragraphs[0].runs:
            # Set font of the headers
            run.font.bold = True
            run.font.size = Pt(header_font_size)

    # Add data from dictionary to the table
    for i, key in enumerate(data.keys()):
        for j, val in enumerate(data[key]):
            # Get the correct cell
            cell = new_table.cell(j + 1, i)
            # Write the text in the cell
            cell.text = str(val)
            text_width = estimate_text_width(str(val), elements_font_size)
            if text_width > max_widths[i]:
                max_widths[i] = text_width
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Set font size of the elements of the table
                    run.font.size = Pt(elements_font_size)

    # Set column widths and apply noWrap to all cells
    for i, width in enumerate(max_widths):
        for cell in new_table.columns[i].cells:
            # Set cell width based on calculated max width
            cell.width = Inches(width)
            # Apply noWrap to each cell
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            no_wrap = OxmlElement('w:noWrap')
            tcPr.append(no_wrap)
            # Apply borders to each cell
            set_cell_borders(cell)

    # Optionally set table borders
    tbl = new_table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    # Define border settings for the table
    border_settings = {
        'w:top': '1',
        'w:left': '1',
        'w:bottom': '1',
        'w:right': '1',
        'w:insideH': '1',
        'w:insideV': '1'
    }

    for border, size in border_settings.items():
        border_elem = OxmlElement(border)
        border_elem.set(qn('w:val'), 'single')
        border_elem.set(qn('w:sz'), size)  # size in half-points
        tblBorders.append(border_elem)

    tblPr.append(tblBorders)


def pdf_to_text(pdf_path, output_txt):
    """
    Extracts text from a PDF file and writes it to a text file.

    Parameters:\n
    - **pdf_path** (``str``): path to the PDF file;
    - **output_txt** (``str``): path to the output txt file.
    """

    # Open the PDF file in read-binary mode
    with open(pdf_path, 'rb') as pdf_file:
        # Create a PdfReader object instead of PdfFileReader
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Initialize an empty string to store the text
        text = ''

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()

    # Write the extracted text to a text file
    with open(output_txt, 'w', encoding='utf-8') as txt_file:
        txt_file.write(text)


def read_text_from_file(file_path: str) -> str:
    """
    Read the text from a file and store it into a string.
    Parameters:\n
    - **file_path** (``str``): absolute path of the txt file to read.
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def write_nice_text(doc: Document, text: str, output_path: str,
                    font_name='Times New Roman', font_size=10, font_color=(0, 0, 0), font_style=None):
    """
    Write a str on a docx.
    Parameters:\n
    - **doc** (``Document``): docx on which write the text;
    - **text** (``str``): text of the paragraph;
    - **font_name** (``str``): name of the font of the paragraph;
    - **font_size** (``int``): size of the title (Pt);
    - **font_color** (``RGBColor``): color of the font;
    - **font_style** (``str``): choose bold, italics or none.
    """
    # Defining a paragraph
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    # Font of the text
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Color of the font
    run.font.color.rgb = RGBColor(*font_color)

    # Style of the font
    if font_style == "bold":
        # Bold text
        run.bold = True
    elif font_style == "italics":
        # Italics text
        run.italics = True
    else:
        pass
    # Set paragraph alignment to left
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Save the file
    doc.save(output_path)


def write_nice_heading(doc: Document, text: str, level: int,
                       font_color: RGBColor, font_name: str, font_size: int, font_style: str):
    """
    Create a nice heading with the following features:\n
    Parameters:\n
    - **doc** (``Document``): docx on which write the heading;
    - **text** (``str``): text of the heading;
    - **level** (``int``): importance level of the heading;
    - **font_color** (``RGBColor``): color of the heading;
    - **font_name** (``str``): name of the font od the text;
    - **style** (``str``): choose bold, italics or none;
    - **font_size** (``int``): size of the heading (Pt).
    """

    # Define Section heading (important for authomatic table of contents)
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]

    # Apply font properties
    run.font.color.rgb = font_color
    run.font.name = font_name
    # Set heading dimension
    run.font.size = Pt(font_size)

    # Apply style
    if font_style == "bold":
        run.bold = True
    elif font_style == "italic":
        run.italic = True

    # Align text to the left
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def extract_voltage_current(bias: list) -> list:
    """
    Helper function to extract values of voltages and currents from a filename with a specific scheme.

    Parameters:\n
    - **bias** (``list``): list of bias names to look for.
    Returns:\n
    - ``list`` of bias values.
    """
    voltage_current_pairs = [re.match(r"(\d+\.\d+)V(\d+\.\d+)mA", b).groups() for b in bias]
    return [vc[0] for vc in voltage_current_pairs], [vc[1] for vc in voltage_current_pairs]


def load_dict_from_filename(filename_path: str) -> dict:
    """
    Parses a filename to extract bias values and LNA names, and returns them in a dictionary.
    The filename must follow a specific pattern to extract the required information:
    - Bias values in the format `Bias(Vg1Ig1_Vg2Ig2_Vg3Ig3)(Vg1Ig1_Vg2Ig2_Vg3Ig3)`
    - LNA names in the format `_Wx..._LNF-...@`

    Parameters:
    - **filename_path** (``str``): path of the file to parse.

    Returns:
    - **result_dict** (``dict``): dict containing the LNA names and corresponding bias values.

    """
    # Regular expression to extract the required parts from the filename
    bias_pattern = r"Bias\((.*?)\)\((.*?)\)"
    lna_pattern = r"_(Wx\..*?)_(LNF-.*?)@"

    # Find all matches using the regex patterns
    bias_match = re.search(bias_pattern, filename_path)
    lna_match = re.search(lna_pattern, filename_path)

    if not (bias_match and lna_match):
        raise ValueError("Filename format is incorrect or doesn't match the expected pattern.")

    # Extract the bias values (coupled string of 'VgIg' per each stage)
    bias1 = bias_match.group(1).split('_')
    bias2 = bias_match.group(2).split('_')

    # Extract LNA names
    lna_names = [lna_match.group(1), lna_match.group(2)]

    # Initialize the dictionary
    result_dict = {
        'LNA name': lna_names,
        'Vd1 [V]': [],
        'Id1 [mA]': [],
        'Vd2 [V]': [],
        'Id2 [mA]': [],
        'Vd3 [V]': [],
        'Id3 [mA]': []
    }

    # Extract and populate the values for Vd1, Id1, Vd2, Id2, Vd3, Id3
    vd1, id1 = extract_voltage_current(bias1)
    vd2, id2 = extract_voltage_current(bias2)

    result_dict['Vd1 [V]'] = [vd1[0], vd2[0]]
    result_dict['Id1 [mA]'] = [id1[0], id2[0]]
    result_dict['Vd2 [V]'] = [vd1[1], vd2[1]]
    result_dict['Id2 [mA]'] = [id1[1], id2[1]]
    result_dict['Vd3 [V]'] = [vd1[2], vd2[2]]
    result_dict['Id3 [mA]'] = [id1[2], id2[2]]

    return result_dict


def set_cell_borders(cell, border_size=1):
    """Set borders for a single cell."""

    # Access the underlying XML element of the cell
    cell_element = cell._element

    # Get or create the cell properties (tcPr) XML element
    cell_pr = cell_element.get_or_add_tcPr()

    # Find the tblBorders element in the cell properties, which defines border settings
    borders = cell_pr.find(qn('w:tblBorders'))

    # If the tblBorders element does not exist, create and append it
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        cell_pr.append(borders)

    # Define border settings for different sides and inner borders
    border_settings = {
        'w:top': border_size,  # Top border
        'w:left': border_size,  # Left border
        'w:bottom': border_size,  # Bottom border
        'w:right': border_size,  # Right border
        'w:insideH': border_size,  # Horizontal inner borders (between cells)
        'w:insideV': border_size  # Vertical inner borders (between cells)
    }

    # Iterate through each border setting
    for border, size in border_settings.items():
        # Find the specific border element within tblBorders
        border_elem = borders.find(qn(border))

        # If the border element does not exist, create and append it
        if border_elem is None:
            border_elem = OxmlElement(border)
            borders.append(border_elem)

        # Set the border style to 'single'
        border_elem.set(qn('w:val'), 'single')

        # Set the border size in half-points
        border_elem.set(qn('w:sz'), str(size))

    return


def set_margins(document, top, bottom, left, right):
    """
    Sets the margins of all sections in a Word document.

    Parameters:\n
    - **document** (``docx.Document``): docx whose margins are to be set.
    - **top** (``float``): size of the top margin in inches.
    - **bottom** (``float``): size of the bottom margin in inches.
    - **left** (``float``): size of the left margin in inches.
    - **right** (``float``): size of the right margin in inches.
    """
    # Collects document sections
    sections = document.sections
    # Set margins
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
